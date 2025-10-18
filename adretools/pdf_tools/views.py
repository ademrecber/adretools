from django.shortcuts import render
from django.http import JsonResponse, HttpResponse
from django.views.decorators.csrf import csrf_exempt
import os
import tempfile
import io


def pdf_home(request):
    tools = [
        {'name': 'Split PDF', 'id': 'split', 'icon': 'fas fa-cut', 'desc': 'Split PDF file into pages'},
        {'name': 'Merge PDF', 'id': 'merge', 'icon': 'fas fa-compress-arrows-alt', 'desc': 'Merge multiple PDFs'},
        {'name': 'Multi Split', 'id': 'multi-split', 'icon': 'fas fa-scissors', 'desc': 'Split specific page ranges'},
        {'name': 'Folder Merge', 'id': 'folder-merge', 'icon': 'fas fa-folder-plus', 'desc': 'Merge same-named PDFs from folders'},
        {'name': 'Interval Split', 'id': 'interval-split', 'icon': 'fas fa-layer-group', 'desc': 'Split by intervals like 3, 5 pages'},
        {'name': 'Convert PDF', 'id': 'convert', 'icon': 'fas fa-exchange-alt', 'desc': 'PDF → Word/Excel/PowerPoint'},
        {'name': 'PDF Security', 'id': 'security', 'icon': 'fas fa-shield-alt', 'desc': 'Encryption, watermark adding'},
        {'name': 'Compress PDF', 'id': 'compress', 'icon': 'fas fa-compress', 'desc': 'Reduce file size'},
        {'name': 'Edit PDF', 'id': 'edit', 'icon': 'fas fa-edit', 'desc': 'Rotate pages, reorder'},
        {'name': 'Image ↔ PDF', 'id': 'image-pdf', 'icon': 'fas fa-images', 'desc': 'Image → PDF, PDF → Image'},
    ]
    return render(request, 'pdf_tools/home.html', {'tools': tools})

@csrf_exempt
def convert_pdf_to_word(request):
    if request.method != 'POST':
        return JsonResponse({'error': 'POST method required'}, status=405)
    
    if 'pdf' not in request.FILES:
        return JsonResponse({'error': 'PDF file required'}, status=400)
    
    pdf_file = request.FILES['pdf']
    
    try:
        from docx import Document
        import PyPDF2
        
        # PDF oku
        reader = PyPDF2.PdfReader(pdf_file)
        
        # Word belgesi oluştur
        doc = Document()
        doc.add_heading(pdf_file.name, 0)
        
        # Her sayfayı işle
        for i, page in enumerate(reader.pages):
            try:
                text = page.extract_text()
                if text and text.strip():
                    doc.add_heading(f'Page {i+1}', 1)
                    doc.add_paragraph(text)
            except:
                doc.add_paragraph(f'[Page {i+1} could not be read]')
        
        # Word dosyasını kaydet
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        
        response = HttpResponse(
            output.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{pdf_file.name.replace(".pdf", ".docx")}"'
        
        return response
        
    except Exception as e:
        return JsonResponse({'error': f'Error: {str(e)}'}, status=500)

@csrf_exempt
def convert_pdf_to_excel(request):
    if request.method != 'POST':
        return JsonResponse({'error': 'POST method required'}, status=405)
    
    if 'pdf' not in request.FILES:
        return JsonResponse({'error': 'PDF file required'}, status=400)
    
    pdf_file = request.FILES['pdf']
    
    try:
        import PyPDF2
        import pandas as pd
        
        # PDF oku
        reader = PyPDF2.PdfReader(pdf_file)
        
        # Tüm metni topla
        all_text = ""
        for page in reader.pages:
            try:
                text = page.extract_text()
                if text:
                    all_text += text + "\n"
            except:
                continue
        
        # Metni satırlara böl ve Excel'e çevir
        lines = [line.strip() for line in all_text.split('\n') if line.strip()]
        
        # Excel için geçersiz karakterleri temizle
        clean_lines = []
        for line in lines:
            # Kontrol karakterlerini kaldır
            clean_line = ''.join(char for char in line if ord(char) >= 32 or char in '\t\n\r')
            clean_lines.append(clean_line[:32767])  # Excel hücre limiti
        
        # DataFrame oluştur
        df = pd.DataFrame({
            'Line No': range(1, len(clean_lines) + 1),
            'Content': clean_lines
        })
        
        # Excel dosyasını kaydet
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name='PDF Content', index=False)
        
        output.seek(0)
        
        response = HttpResponse(
            output.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
        response['Content-Disposition'] = f'attachment; filename="{pdf_file.name.replace(".pdf", ".xlsx")}"'
        
        return response
        
    except Exception as e:
        return JsonResponse({'error': f'Error: {str(e)}'}, status=500)

@csrf_exempt
def convert_word_to_pdf(request):
    return JsonResponse({
        'error': 'Word → PDF conversion is currently under development. You can use other features!'
    }, status=501)

@csrf_exempt
def split_pdf(request):
    if request.method != 'POST':
        return JsonResponse({'error': 'POST method required'}, status=405)
    
    if 'pdf' not in request.FILES:
        return JsonResponse({'error': 'PDF file required'}, status=400)
    
    pdf_file = request.FILES['pdf']
    
    try:
        from PyPDF2 import PdfReader, PdfWriter
        import zipfile
        
        # PDF'i oku
        reader = PdfReader(pdf_file)
        total_pages = len(reader.pages)
        
        # ZIP dosyası oluştur
        zip_buffer = io.BytesIO()
        
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            # Her sayfayı ayrı PDF olarak kaydet
            for i, page in enumerate(reader.pages):
                writer = PdfWriter()
                writer.add_page(page)
                
                # Sayfa PDF'ini oluştur
                page_buffer = io.BytesIO()
                writer.write(page_buffer)
                page_buffer.seek(0)
                
                # ZIP'e ekle
                filename = f"{pdf_file.name.replace('.pdf', '')}_sayfa_{i+1}.pdf"
                zip_file.writestr(filename, page_buffer.getvalue())
        
        zip_buffer.seek(0)
        
        response = HttpResponse(zip_buffer.getvalue(), content_type='application/zip')
        response['Content-Disposition'] = f'attachment; filename="{pdf_file.name.replace(".pdf", "_bolunmus.zip")}"'
        
        return response
        
    except Exception as e:
        return JsonResponse({'error': f'Error: {str(e)}'}, status=500)

@csrf_exempt
def merge_pdf(request):
    if request.method != 'POST':
        return JsonResponse({'error': 'POST method required'}, status=405)
    
    pdf_files = request.FILES.getlist('pdfs')
    
    if len(pdf_files) < 2:
        return JsonResponse({'error': 'At least 2 PDF files required'}, status=400)
    
    try:
        from PyPDF2 import PdfReader, PdfWriter
        
        writer = PdfWriter()
        
        # Tüm PDF'leri birleştir
        for pdf_file in pdf_files:
            reader = PdfReader(pdf_file)
            for page in reader.pages:
                writer.add_page(page)
        
        # Birleştirilmiş PDF'i kaydet
        output_buffer = io.BytesIO()
        writer.write(output_buffer)
        output_buffer.seek(0)
        
        response = HttpResponse(output_buffer.getvalue(), content_type='application/pdf')
        response['Content-Disposition'] = 'attachment; filename="merged_pdf.pdf"'
        
        return response
        
    except Exception as e:
        return JsonResponse({'error': f'Error: {str(e)}'}, status=500)

@csrf_exempt
def compress_pdf(request):
    if request.method != 'POST':
        return JsonResponse({'error': 'POST method required'}, status=405)
    
    if 'pdf' not in request.FILES:
        return JsonResponse({'error': 'PDF file required'}, status=400)
    
    pdf_file = request.FILES['pdf']
    
    try:
        from PyPDF2 import PdfReader, PdfWriter
        
        reader = PdfReader(pdf_file)
        writer = PdfWriter()
        
        # Sayfaları kopyala ve sıkıştır
        for page in reader.pages:
            page.compress_content_streams()
            writer.add_page(page)
        
        # Sıkıştırılmış PDF'i kaydet
        output_buffer = io.BytesIO()
        writer.write(output_buffer)
        output_buffer.seek(0)
        
        response = HttpResponse(output_buffer.getvalue(), content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="{pdf_file.name.replace(".pdf", "_compressed.pdf")}"'
        
        return response
        
    except Exception as e:
        return JsonResponse({'error': f'Error: {str(e)}'}, status=500)

@csrf_exempt
def encrypt_pdf(request):
    if request.method != 'POST':
        return JsonResponse({'error': 'POST method required'}, status=405)
    
    if 'pdf' not in request.FILES or 'password' not in request.POST:
        return JsonResponse({'error': 'PDF file and password required'}, status=400)
    
    pdf_file = request.FILES['pdf']
    password = request.POST['password']
    
    try:
        from PyPDF2 import PdfReader, PdfWriter
        
        # Geçici dosya oluştur
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_pdf:
            for chunk in pdf_file.chunks():
                temp_pdf.write(chunk)
            temp_pdf_path = temp_pdf.name
        
        # PDF'i oku
        reader = PdfReader(temp_pdf_path)
        writer = PdfWriter()
        
        # Tüm sayfaları kopyala
        for page in reader.pages:
            writer.add_page(page)
        
        # Şifrele
        writer.encrypt(password)
        
        # Şifrelenmiş PDF'i kaydet
        output_buffer = io.BytesIO()
        writer.write(output_buffer)
        output_buffer.seek(0)
        
        # Geçici dosyayı sil
        os.unlink(temp_pdf_path)
        
        response = HttpResponse(output_buffer.getvalue(), content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="{pdf_file.name.replace(".pdf", "_sifreli.pdf")}"'
        
        return response
        
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

@csrf_exempt
def add_watermark(request):
    if request.method != 'POST':
        return JsonResponse({'error': 'POST method required'}, status=405)
    
    if 'pdf' not in request.FILES or 'watermark_text' not in request.POST:
        return JsonResponse({'error': 'PDF file and watermark text required'}, status=400)
    
    pdf_file = request.FILES['pdf']
    watermark_text = request.POST['watermark_text']
    position = request.POST.get('position', 'center')
    opacity = float(request.POST.get('opacity', '0.5'))
    font_size = int(request.POST.get('size', '36'))
    
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        from PyPDF2 import PdfReader, PdfWriter
        
        # Geçici dosya oluştur
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_pdf:
            for chunk in pdf_file.chunks():
                temp_pdf.write(chunk)
            temp_pdf_path = temp_pdf.name
        
        # PDF'i oku
        reader = PdfReader(temp_pdf_path)
        writer = PdfWriter()
        
        # Her sayfa için filigran ekle
        for page_num, page in enumerate(reader.pages):
            # Filigran PDF'i oluştur
            watermark_buffer = io.BytesIO()
            
            # Sayfa boyutlarını al
            page_width = float(page.mediabox.width)
            page_height = float(page.mediabox.height)
            
            # Filigran canvas'ı oluştur
            c = canvas.Canvas(watermark_buffer, pagesize=(page_width, page_height))
            
            # Pozisyona göre koordinatları belirle
            if position == 'center':
                x, y = page_width / 2, page_height / 2
            elif position == 'top-left':
                x, y = 50, page_height - 50
            elif position == 'top-right':
                x, y = page_width - 50, page_height - 50
            elif position == 'bottom-left':
                x, y = 50, 50
            elif position == 'bottom-right':
                x, y = page_width - 50, 50
            else:
                x, y = page_width / 2, page_height / 2
            
            # Filigran metnini ekle
            c.setFillColorRGB(0.5, 0.5, 0.5, alpha=opacity)
            c.setFont("Helvetica-Bold", font_size)
            
            # Metni döndür (45 derece)
            c.saveState()
            c.translate(x, y)
            c.rotate(45)
            # Metni ortala
            text_width = c.stringWidth(watermark_text, "Helvetica-Bold", font_size)
            c.drawString(-text_width/2, -font_size/2, watermark_text)
            c.restoreState()
            
            c.save()
            watermark_buffer.seek(0)
            
            # Filigran PDF'ini oku
            watermark_reader = PdfReader(watermark_buffer)
            watermark_page = watermark_reader.pages[0]
            
            # Ana sayfaya filigranı ekle
            page.merge_page(watermark_page)
            writer.add_page(page)
        
        # Filigranli PDF'i kaydet
        output_buffer = io.BytesIO()
        writer.write(output_buffer)
        output_buffer.seek(0)
        
        # Geçici dosyayı sil
        os.unlink(temp_pdf_path)
        
        response = HttpResponse(output_buffer.getvalue(), content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="{pdf_file.name.replace(".pdf", "_filigranli.pdf")}"'
        
        return response
        
    except Exception as e:
        return JsonResponse({'error': f'Filigran ekleme hatası: {str(e)}'}, status=500)